"""
Ultrasound Report Generator (DOCX)
====================================
Generates patient-specific ultrasound reports as editable .docx files.
All standard/boilerplate text is pre-filled. Doctor adds impression in Word.

Supported report types:
  1  - Early Pregnancy Scan
  2  - NT Scan (Nuchal Translucency)
  3  - Anomaly Scan
  4  - Growth Scan
  5  - Follicular Study
  6  - Abdomen & Pelvis (Female)
  7  - Abdomen & Pelvis (Male)

Usage (interactive):
    python fill_report.py

Usage (programmatic):
    from fill_report import generate_report
    generate_report("early_pregnancy", {
        "patient_name": "Vidya Shree",
        "age": "24",
        "date": "01/07/2026",
        ...
    })
"""

from __future__ import annotations

import argparse
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DOCTOR_NAME = "Dr. Suhas.R.H."
DOCTOR_QUAL = "MBBS. MDRD."
REFERRING_DEFAULT = "Dr. Latha.K.P."
OUTPUT_DIR = Path("reports/filled")

STANDARD_NOTE = (
    "Note: This is only a radiological impression and not a diagnosis and has its "
    "limitations. Therefore, it should be interpreted in correlation with clinical "
    "and/or pathological findings."
)

NT_NOTE = (
    "Note: foetal heart, spine, face and kidneys cannot be well assessed at this "
    "stage of pregnancy. These regions can be assessed in the 20 to 24 weeks scan.\n\n"
    "Note:\nDown's syndrome can't be diagnosed on the base of USG alone.\n"
    "The risk ratio indicates risk rates – it's not definitive testing.\n"
    "The detection of downs syndrome by:\n"
    "  First trimester NT only – 64 to 70%.\n"
    "  First trimester combined (NT + maternal blood test) 80–85%.\n"
    "  Sequential screening (combined quadruple 15–19 weeks + genetic sonogram "
    "at 18–20 weeks) 95%.\n"
    "  Maternal blood test for cell free foetal DNA 99%.\n"
    "  Invasive testing (CVS/amniocentesis) – procedure related risk of about 1:200.\n"
    "This has been explained to the patient and attendant."
)

ANOMALY_NOTE = (
    "Note:\n"
    "Down's syndrome can't be diagnosed on the base of USG alone.\n"
    "Sequential screening (combined quadruple 15–19 weeks + genetic sonogram at 18–20 weeks) 95%.\n"
    "Maternal blood test for cell free foetal DNA 99%.\n"
    "Invasive testing (CVS/amniocentesis) – procedure related risk of about 1:200.\n"
    "This has been explained to the patient and attendant.\n\n"
    "# Ultrasound can't detect all congenital anomalies. Detection rate of congenital "
    "anomalies by anomaly scan is 60 to 80%.\n\n"
    "Please note:\nAll abnormalities and genetic syndromes cannot be ruled out by "
    "ultrasound examination. Ultrasound has its own limitations. Some abnormalities "
    "evolve as the gestation advances. Detection rate depends on gestational age, "
    "foetal position, tissue penetration of sound waves, and patient body habits."
)

GROWTH_NOTE = (
    "Note: foetal anatomy (limbs, spine, cardiac, etc.) can't be made out in detail "
    "at this stage due to advanced gestational age and un-accommodating foetal position. "
    "The ideal time to study foetal anatomy is 20 to 22 weeks. "
    "This has been explained to the patient and attendant.\n"
    "# All congenital anomalies can't be detected by antenatal ultrasound."
)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _new_doc() -> Document:
    """Create a document with standard A4 margins."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    # Remove default paragraph spacing
    style = doc.styles["Normal"]
    style.paragraph_format.space_after = Pt(2)
    return doc


def _heading(doc: Document, text: str, size: int = 13, center: bool = True) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _line(doc: Document, *parts: tuple[str, str, bool]) -> None:
    """
    Add a paragraph with label/value pairs.
    parts: list of (label, value, bold_label)
    """
    p = doc.add_paragraph()
    for label, value, bold in parts:
        if label:
            r = p.add_run(label)
            r.bold = bold
        if value:
            p.add_run(value)


def _section_label(doc: Document, label: str, value: str = "") -> None:
    """Bold label: normal value on one line."""
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    if value:
        p.add_run(value)


def _plain(doc: Document, text: str, indent: bool = False) -> None:
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)


def _patient_header(doc: Document, patient_name: str, age: str, date: str) -> None:
    """Bold-label patient info line: Patient Name / Age / Date."""
    p = doc.add_paragraph()
    p.add_run("Patient Name: ").bold = True
    p.add_run(f"{patient_name}          ")
    p.add_run("Age: ").bold = True
    p.add_run(f"{age}          ")
    p.add_run("Date: ").bold = True
    p.add_run(date)


def _small_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)


def _doppler_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Add the Doppler PI index table."""
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Artery"
    hdr[1].text = "PI"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for i, (artery, pi) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = artery
        table.rows[i].cells[1].text = pi


def _declaration(doc: Document, patient_name: str, sex: str = "F") -> None:
    """sex: 'F' → her/Mrs, 'M' → his/Mr, anything else → their/the patient."""
    if sex == "M":
        pronoun, title = "his", "Mr."
    elif sex == "F":
        pronoun, title = "her", "Mrs."
    else:
        pronoun, title = "their", ""
    name_with_title = f"{title} {patient_name}".strip()
    doc.add_paragraph()
    p = doc.add_paragraph(
        "DECLARATION OF DOCTOR/PERSON CONDUCTING ULTRASONOGRAPHY / IMAGE SCANNING"
    )
    for run in p.runs:
        run.bold = True
    doc.add_paragraph(
        f"I. {DOCTOR_NAME} declare that while conducting ultrasonography / image "
        f"scanning on {name_with_title}, I have neither detected nor disclosed the "
        f"sex of {pronoun} foetus to anybody in any manner."
    )


def _doctor_signature(doc: Document, right_align: bool = True) -> None:
    p = doc.add_paragraph()
    if right_align:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"{DOCTOR_NAME}\n{DOCTOR_QUAL}").bold = False


def _save(doc: Document, patient_name: str, scan_type: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(c for c in patient_name if c.isalnum() or c in " ._-").strip()
    date_str = datetime.today().strftime("%Y%m%d")
    filename = f"{safe_name}_{scan_type}_{date_str}.docx"
    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# 1. Early Pregnancy Scan
# ---------------------------------------------------------------------------

def generate_early_pregnancy_report(data: dict) -> Path:
    """
    Required keys: patient_name, age, date, ref_by, lmp
    Optional keys: clinical_details, uterus_notes, crl, ga_scan_weeks, ga_scan_days,
                   ga_lmp_weeks, ga_lmp_days, edd_scan, edd_lmp, fhr,
                   cervix_length, ovaries_notes, impression
    """
    d = data
    doc = _new_doc()

    _heading(doc, "EARLY PREGNANCY SCAN")
    _patient_header(doc, d.get('patient_name', ''), d.get('age', ''), d.get('date', ''))
    _section_label(doc, "Ref by", d.get("ref_by", REFERRING_DEFAULT))
    _section_label(doc, "Clinical details / diagnosis", d.get("clinical_details", ""))
    _section_label(doc, "LMP", d.get("lmp", ""))

    doc.add_paragraph()
    _section_label(doc, "Uterus",
                   d.get("uterus_notes", "Gravid uterus."))
    _section_label(doc, "Gestational Sac", "Seen")

    p = doc.add_paragraph()
    p.add_run("Embryo: ").bold = True
    p.add_run("Single in number")

    # CRL / GA / EDD
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.4)
    p2.add_run("CRL: ").bold = True
    p2.add_run(f"{d.get('crl', '')} cms")

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Inches(0.4)
    p3.add_run("GA: ").bold = True
    p3.add_run(
        f"{d.get('ga_scan_weeks', '')} weeks {d.get('ga_scan_days', '')} days"
        f"                    "
    )
    p3.add_run("GA by LMP: ").bold = True
    p3.add_run(f"{d.get('ga_lmp_weeks', '')} weeks {d.get('ga_lmp_days', '')} days")

    p4 = doc.add_paragraph()
    p4.paragraph_format.left_indent = Inches(0.4)
    p4.add_run("EDD: ").bold = True
    p4.add_run(f"{d.get('edd_scan', '')}                    ")
    p4.add_run("EDD by LMP: ").bold = True
    p4.add_run(f"{d.get('edd_lmp', '')}")

    _section_label(doc, "Cardiac activity", f"GOOD ({d.get('fhr', '')} BPM)")
    _section_label(doc, "Cervix",
                   f"Length {d.get('cervix_length', '')} cms.  Internal os closed.")
    _section_label(doc, "Ovaries & adnexa", d.get("ovaries_notes", ""))

    doc.add_paragraph()
    _section_label(doc, "Impression", "")
    imp_para = doc.add_paragraph(d.get("impression", ""))
    imp_para.paragraph_format.left_indent = Inches(0.2)

    doc.add_paragraph()
    _declaration(doc, d.get("patient_name", ""), sex=d.get("_sex", "F"))
    _doctor_signature(doc)
    return _save(doc, d.get("patient_name", "patient"), "early_pregnancy")


# ---------------------------------------------------------------------------
# 2. NT Scan
# ---------------------------------------------------------------------------

def generate_nt_scan_report(data: dict) -> Path:
    """
    Required keys: patient_name, age, date, ref_by, lmp
    Optional keys: clinical_details, ga_lmp_weeks, ga_lmp_days, edd_lmp,
                   crl, aua_weeks, aua_days, fhr, edd_scan,
                   nasal_bone, nt, ductus_venosus,
                   liquor, placenta_location, placenta_grade, cervix_length,
                   ovaries_notes, impression,
                   doppler_right_pi, doppler_left_pi
    """
    d = data
    doc = _new_doc()

    _heading(doc, "OBSTETRIC ULTRASOUND REPORT")
    _patient_header(doc, d.get('patient_name', ''), d.get('age', ''), d.get('date', ''))
    _section_label(doc, "Ref by", d.get("ref_by", REFERRING_DEFAULT))
    _section_label(doc, "Clinical details / diagnosis",
                   d.get("clinical_details", "For early pregnancy assessment & NT scan."))
    doc.add_paragraph(
        f"LMP: {d.get('lmp', '')}          "
        f"GA from LMP: {d.get('ga_lmp_weeks', '')} weeks {d.get('ga_lmp_days', '')} days"
        f"          EDD: {d.get('edd_lmp', '')}"
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Foetus: ").bold = True
    p.add_run("Single live foetus")

    _heading(doc, "Biometric Measurements", size=11, center=False)
    doc.add_paragraph(
        f"CRL: {d.get('crl', '')} cms"
    )
    doc.add_paragraph(
        f"AUA: {d.get('aua_weeks', '')} weeks {d.get('aua_days', '')} days          "
        f"FHR: {d.get('fhr', '')} BPM"
    )
    doc.add_paragraph(f"EDD: {d.get('edd_scan', '')}")

    doc.add_paragraph()
    _heading(doc, "ANEUPLOIDY MARKERS", size=11, center=False)
    _section_label(doc, "Nasal bone", f"{d.get('nasal_bone', '')} mm")
    _section_label(doc, "Nuchal translucency", f"{d.get('nt', '')} mm")
    _section_label(doc, "Ductus venosus flow",
                   d.get("ductus_venosus", ""))

    doc.add_paragraph()
    _section_label(doc, "LIQUOR", d.get("liquor", ""))
    _section_label(
        doc, "PLACENTA",
        f"Located in {d.get('placenta_location', 'fundal posterior')} "
        f"grade {d.get('placenta_grade', '1')}, lower limit of placenta seen "
        f"well away from the internal OS."
    )
    _section_label(
        doc, "CERVIX",
        f"Length {d.get('cervix_length', '')} cms, internal OS closed."
    )
    doc.add_paragraph(
        "Uterine artery doppler study shows, recorded indices are as follows."
    )
    _doppler_table(doc, [
        ("Right uterine artery", d.get("doppler_right_pi", "")),
        ("Left uterine artery", d.get("doppler_left_pi", "")),
    ])
    _section_label(doc, "Ovaries & adnexa", d.get("ovaries_notes", ""))

    doc.add_paragraph()
    _heading(doc, "IMPRESSION", size=11, center=False)
    doc.add_paragraph(d.get("impression", ""))

    doc.add_paragraph()
    _small_note(doc, NT_NOTE)
    _declaration(doc, d.get("patient_name", ""), sex=d.get("_sex", "F"))
    _doctor_signature(doc)
    return _save(doc, d.get("patient_name", "patient"), "nt_scan")


# ---------------------------------------------------------------------------
# 3. Anomaly Scan
# ---------------------------------------------------------------------------

def generate_anomaly_scan_report(data: dict) -> Path:
    """
    Required keys: patient_name, age, date, ref_by, lmp
    Optional keys: clinical_details, bpd, hc, fl, ac, hl, tl, rl, fib, ul,
                   aua_weeks, aua_days, ga_lmp_weeks, ga_lmp_days,
                   edd_scan, edd_lmp, efw, efw_error, fl_ac_ratio,
                   nasal_bone_length, nuchal_fold, tcd, cisterna_magna,
                   lvta, foot_length, fhr, foetal_lie,
                   head_notes, face_notes, thorax_notes, heart_notes,
                   abdomen_notes, limbs_notes, spine_notes,
                   liquor, placenta_location, placenta_grade,
                   cervix_length, cord_notes,
                   impression, comments,
                   doppler_right_pi, doppler_left_pi
    """
    d = data
    doc = _new_doc()

    _heading(doc, "ANOMALY SCAN")
    _patient_header(doc, d.get('patient_name', ''), d.get('age', ''), d.get('date', ''))
    doc.add_paragraph(
        f"Ref By: {d.get('ref_by', REFERRING_DEFAULT)}          "
        f"LMP: {d.get('lmp', '')}"
    )
    _section_label(doc, "Clinical details / diagnosis", d.get("clinical_details", ""))

    p = doc.add_paragraph()
    p.add_run("Foetus: ").bold = True
    p.add_run(f"Single live foetus with {d.get('foetal_lie', 'cephalic')} presentation.")

    doc.add_paragraph()
    _heading(doc, "Biometric Measurements", size=11, center=False)
    doc.add_paragraph(
        f"BPD: {d.get('bpd', '')} cms          HC: {d.get('hc', '')} cms"
    )
    doc.add_paragraph(
        f"FL:  {d.get('fl', '')} cms          AC: {d.get('ac', '')} cms"
    )
    doc.add_paragraph(
        f"HL:  {d.get('hl', '')} cms          TL: {d.get('tl', '')} cms"
    )
    doc.add_paragraph(
        f"RL:  {d.get('rl', '')} cms          FIB: {d.get('fib', '')} cms"
    )
    doc.add_paragraph(f"UL:  {d.get('ul', '')} cms")
    doc.add_paragraph(
        f"AUA: {d.get('aua_weeks', '')} weeks {d.get('aua_days', '')} days          "
        f"GA by LMP: {d.get('ga_lmp_weeks', '')} weeks {d.get('ga_lmp_days', '')} days"
    )
    doc.add_paragraph(
        f"EDD: {d.get('edd_scan', '')}          EDD by LMP: {d.get('edd_lmp', '')}"
    )
    doc.add_paragraph(
        f"EFW: {d.get('efw', '')} gms +/- {d.get('efw_error', '')} gms          "
        f"FL/AC: {d.get('fl_ac_ratio', '')} %"
    )
    doc.add_paragraph(
        f"Nasal Bone Length: {d.get('nasal_bone_length', '')} mm          "
        f"Nuchal Fold Thickness: {d.get('nuchal_fold', '')} mm"
    )
    doc.add_paragraph(
        f"Trans cerebellar diameter: {d.get('tcd', '')} cm          "
        f"Cisterna magna: {d.get('cisterna_magna', '')} cm"
    )
    doc.add_paragraph(
        f"Transverse dimension of lateral ventricular atrium: {d.get('lvta', '')} cm          "
        f"Foot length: {d.get('foot_length', '')} cms"
    )
    doc.add_paragraph(
        f"Foetal cardiac activity and movements good.  FHR: {d.get('fhr', '')} BPM."
    )

    # Organ findings
    doc.add_paragraph()
    _heading(doc, "Foetal head:", size=11, center=False)
    doc.add_paragraph(d.get("head_notes", ""))

    _heading(doc, "Foetal face:", size=11, center=False)
    doc.add_paragraph(d.get("face_notes", ""))

    _heading(doc, "Foetal thorax:", size=11, center=False)
    doc.add_paragraph(d.get("thorax_notes", "Both lungs seen."))

    _heading(doc, "Foetal heart:", size=11, center=False)
    doc.add_paragraph(d.get("heart_notes", ""))

    _heading(doc, "Foetal abdomen:", size=11, center=False)
    doc.add_paragraph(d.get("abdomen_notes", ""))

    _heading(doc, "Foetal limbs:", size=11, center=False)
    doc.add_paragraph(d.get("limbs_notes", ""))

    _heading(doc, "Foetal spine:", size=11, center=False)
    doc.add_paragraph(d.get("spine_notes", ""))

    doc.add_paragraph()
    _section_label(doc, "Liquor", d.get("liquor", ""))
    _section_label(
        doc, "Placenta",
        f"Located in {d.get('placenta_location', '')} "
        f"grade {d.get('placenta_grade', '')}. "
        "Lower limit of placenta situated well away from the internal OS."
    )
    _section_label(
        doc, "Cervix",
        f"Length {d.get('cervix_length', '')} cms, internal OS closed."
    )
    _section_label(
        doc, "Umbilical cord",
        d.get("cord_notes", "")
    )
    doc.add_paragraph(
        "Uterine artery doppler study shows, recorded indices are as follows."
    )
    _doppler_table(doc, [
        ("Right uterine artery", d.get("doppler_right_pi", "")),
        ("Left uterine artery", d.get("doppler_left_pi", "")),
    ])

    doc.add_paragraph()
    _heading(doc, "IMPRESSION", size=11, center=False)
    doc.add_paragraph(d.get("impression", ""))
    if d.get("comments"):
        doc.add_paragraph()
        _section_label(doc, "Comments", "")
        doc.add_paragraph(d["comments"])

    doc.add_paragraph()
    _small_note(doc, ANOMALY_NOTE)
    _declaration(doc, d.get("patient_name", ""), sex=d.get("_sex", "F"))
    _doctor_signature(doc)
    return _save(doc, d.get("patient_name", "patient"), "anomaly_scan")


# ---------------------------------------------------------------------------
# 4. Growth Scan
# ---------------------------------------------------------------------------

def generate_growth_scan_report(data: dict) -> Path:
    """
    Required keys: patient_name, age, date, ref_by, lmp
    Optional keys: clinical_details, foetal_presentation, edd_lmp, edd_scan,
                   bpd, hc, fl, ac, aua_weeks, aua_days, ga_lmp_weeks, ga_lmp_days,
                   efw, efw_error, fhr, liquor, afi, cpr,
                   placenta_location, placenta_grade, cervix_length,
                   impression,
                   doppler_right_pi, doppler_left_pi, doppler_umbilical_pi, doppler_mca_pi
    """
    d = data
    doc = _new_doc()

    _patient_header(doc, d.get('patient_name', ''), d.get('age', ''), d.get('date', ''))
    _section_label(doc, "Ref By", d.get("ref_by", REFERRING_DEFAULT))

    _heading(doc, "OBSTETRIC ULTRASOUND REPORT")
    _section_label(doc, "Clinical details / diagnosis", d.get("clinical_details", ""))
    _section_label(doc, "LMP", d.get("lmp", ""))
    doc.add_paragraph(
        "Type of ultrasound study: Growth assessment and biophysical profile only."
    )
    doc.add_paragraph(
        f"Route: Trans abdominal          "
        f"EDD by LMP: {d.get('edd_lmp', '')}          "
        f"EDD by AUA: {d.get('edd_scan', '')}"
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Foetus: ").bold = True
    p.add_run(
        f"Single live foetus with {d.get('foetal_presentation', 'cephalic')} presentation."
    )

    _heading(doc, "Biometric Measurements:", size=11, center=False)
    doc.add_paragraph(
        f"BPD: {d.get('bpd', '')} cms          HC: {d.get('hc', '')} cms"
    )
    doc.add_paragraph(
        f"FL:  {d.get('fl', '')} cms          AC: {d.get('ac', '')} cms"
    )
    doc.add_paragraph(
        f"AUA: {d.get('aua_weeks', '')} weeks {d.get('aua_days', '')} days          "
        f"GA by LMP: {d.get('ga_lmp_weeks', '')} weeks {d.get('ga_lmp_days', '')} days"
    )
    doc.add_paragraph(
        f"EFW: {d.get('efw', '')} gms  +/-  {d.get('efw_error', '')} gms"
    )

    doc.add_paragraph()
    _heading(doc, "Biophysical profile:", size=11, center=False)
    doc.add_paragraph(
        "Foetal maturity – femoral ossification centre noted and tibial appearing.\n"
        "Foetal tone and movements are good.\n"
        "Foetal breathing movements good.\n"
        "U/S biophysical scoring 8/8.\n"
        "Foetal cardiac activity good."
        f"          FHR: {d.get('fhr', '')} BPM."
    )
    doc.add_paragraph(
        f"Liquor: {d.get('liquor', '')}          AFI: {d.get('afi', '')} cms."
    )
    if d.get("cpr"):
        _section_label(doc, "CPR", d["cpr"])

    doc.add_paragraph()
    _section_label(
        doc, "Placenta",
        f"Located in {d.get('placenta_location', 'posterior')}, "
        f"grade {d.get('placenta_grade', '3')}. "
        "Lower limit of placenta situated well away from the internal OS."
    )
    _section_label(
        doc, "Cervix",
        f"Visualized cervix measuring {d.get('cervix_length', '')} cms."
    )
    doc.add_paragraph(
        "Umbilical cord: No E/O cord around the neck noted at the time of scanning."
    )

    doc.add_paragraph()
    _heading(doc, "IMPRESSION", size=11, center=False)
    doc.add_paragraph(d.get("impression", ""))

    doc.add_paragraph()
    _small_note(doc, GROWTH_NOTE)
    _declaration(doc, d.get("patient_name", ""), sex=d.get("_sex", "F"))
    doc.add_paragraph("Uterine artery doppler study – recorded indices:")
    _doppler_table(doc, [
        ("Right uterine artery", d.get("doppler_right_pi", "")),
        ("Left uterine artery", d.get("doppler_left_pi", "")),
        ("Umbilical artery", d.get("doppler_umbilical_pi", "")),
        ("MCA flow", d.get("doppler_mca_pi", "")),
    ])
    _doctor_signature(doc)
    return _save(doc, d.get("patient_name", "patient"), "growth_scan")


# ---------------------------------------------------------------------------
# 5. Follicular Study
# ---------------------------------------------------------------------------

def generate_follicular_study_report(data: dict) -> Path:
    """
    Required keys: patient_name, age, date, ref_by
    Optional keys: clinical_details, uterus_size, right_ovary_size, left_ovary_size,
                   bilateral_notes, impression,
                   follicle_rows (list of dicts with keys:
                     day, endometrial_thickness, right_follicle, left_follicle, free_fluid)
    """
    d = data
    doc = _new_doc()

    _heading(doc, "PELVIC ULTRASOUND REPORT (FOLLICULAR STUDY)")
    _section_label(doc, "Clinical data",
                   d.get("clinical_details", "c/o anxious to conceive."))
    _section_label(doc, "URINARY BLADDER", "Distended.")
    _section_label(doc, "UTERUS",
                   f"Anteverted, Measuring {d.get('uterus_size', '')} cms.")
    _section_label(doc, "CERVIX", "Normal.")
    _section_label(doc, "RIGHT OVARY",
                   f"Measuring {d.get('right_ovary_size', '')} cms.")
    _section_label(doc, "LEFT OVARY",
                   f"Measuring {d.get('left_ovary_size', '')} cms.")

    if d.get("bilateral_notes"):
        doc.add_paragraph(d["bilateral_notes"])

    doc.add_paragraph()
    _section_label(doc, "Impression", "")
    doc.add_paragraph(d.get("impression", ""))

    doc.add_paragraph()
    _small_note(doc, STANDARD_NOTE)
    _doctor_signature(doc)

    # Patient info block
    doc.add_paragraph()
    doc.add_paragraph(f"NAME: {d.get('patient_name', '')}")
    doc.add_paragraph(
        f"REFERENCE – {d.get('ref_by', REFERRING_DEFAULT)}          "
        f"DATE: {d.get('date', '')}"
    )
    doc.add_paragraph(f"AGE/SEX: {d.get('age', '')} Y/F")

    doc.add_paragraph()
    # Follicular tracking table
    headers = [
        "DAY",
        "ENDOMETRIAL\nTHICKNESS",
        "ANTRAL FOLLICLE\n(RIGHT OVARY)",
        "ANTRAL FOLLICLE\n(LEFT OVARY)",
        "FREE\nFLUID",
    ]
    follicle_rows = d.get("follicle_rows", [{}] * 12)
    table = doc.add_table(rows=1 + len(follicle_rows), cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for i, row_data in enumerate(follicle_rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = str(row_data.get("day", ""))
        cells[1].text = str(row_data.get("endometrial_thickness", ""))
        cells[2].text = str(row_data.get("right_follicle", ""))
        cells[3].text = str(row_data.get("left_follicle", ""))
        cells[4].text = str(row_data.get("free_fluid", ""))

    _doctor_signature(doc)
    return _save(doc, d.get("patient_name", "patient"), "follicular_study")


# ---------------------------------------------------------------------------
# 6. Abdomen & Pelvis – Female
# ---------------------------------------------------------------------------

def generate_abdomen_pelvis_female_report(data: dict) -> Path:
    """
    Required keys: patient_name, age, date, ref_by
    Optional keys: clinical_data,
                   liver_size, gallbladder_notes, pancreas_notes, spleen_size,
                   right_kidney_size, left_kidney_size,
                   uterus_size, endometrium_mm, right_ovary_size, left_ovary_size,
                   free_fluid_notes, bowel_notes, impression
    """
    d = data
    doc = _new_doc()

    _heading(doc, "ABDOMINO-PELVIC ULTRASOUND REPORT")
    _patient_header(doc, d.get("patient_name", ""), d.get("age", ""), d.get("date", ""))
    _section_label(doc, "Ref by", d.get("ref_by", REFERRING_DEFAULT))
    _section_label(doc, "Clinical data", d.get("clinical_data", ""))

    _section_label(
        doc, "LIVER",
        f"Measuring {d.get('liver_size', '')} cms."


    )
    _section_label(
        doc, "GALL BLADDER",
        d.get("gallbladder_notes", "")
    )
    _section_label(
        doc, "PANCREAS",
        d.get("pancreas_notes", "")
    )
    _section_label(
        doc, "SPLEEN",
        f"Spleen measuring {d.get('spleen_size', '')} cms."
    )
    _section_label(
        doc, "RIGHT KIDNEY",
        f"Measuring {d.get('right_kidney_size', '')} cms."
    )
    _section_label(
        doc, "LEFT KIDNEY",
        f"Measuring {d.get('left_kidney_size', '')} cms."
    )
    _section_label(doc, "URINARY BLADDER", "Distended.")
    _section_label(
        doc, "UTERUS",
        f"Anteverted, measuring {d.get('uterus_size', '')} cms, "
        f"endometrium thickness {d.get('endometrium_mm', '')} mm."
    )
    _section_label(
        doc, "RIGHT OVARY", f"Measuring {d.get('right_ovary_size', '')} cms."
    )
    _section_label(
        doc, "LEFT OVARY", f"Measuring {d.get('left_ovary_size', '')} cms."
    )
    _section_label(
        doc, "FREE FLUID",
        d.get("free_fluid_notes", "")
    )
    _section_label(
        doc, "BOWEL LOOPS",
        d.get("bowel_notes", "Visualized bowel loops show peristalsis.")
    )

    doc.add_paragraph()
    _heading(doc, "IMPRESSION:", size=11, center=False)
    doc.add_paragraph(d.get("impression", ""))

    doc.add_paragraph()
    _small_note(doc, STANDARD_NOTE)
    _doctor_signature(doc)

    return _save(doc, d.get("patient_name", "patient"), "abdomen_pelvis_female")


# ---------------------------------------------------------------------------
# 7. Abdomen & Pelvis – Male
# ---------------------------------------------------------------------------

def generate_abdomen_pelvis_male_report(data: dict) -> Path:
    """
    Required keys: patient_name, age, date, ref_by
    Optional keys: clinical_data,
                   liver_size, gallbladder_notes, pancreas_notes, spleen_size,
                   right_kidney_size, left_kidney_size,
                   urinary_bladder_notes, prostate_notes,
                   free_fluid_notes, bowel_notes, impression
    """
    d = data
    doc = _new_doc()

    _heading(doc, "ABDOMINO-PELVIC ULTRASOUND REPORT")
    _patient_header(doc, d.get("patient_name", ""), d.get("age", ""), d.get("date", ""))
    _section_label(doc, "Ref by", d.get("ref_by", REFERRING_DEFAULT))
    _section_label(doc, "Clinical data", d.get("clinical_data", ""))

    _section_label(
        doc, "LIVER",
        f"Measuring {d.get('liver_size', '')} cms."


    )
    _section_label(
        doc, "GALL BLADDER",
        d.get("gallbladder_notes", "")
    )
    _section_label(
        doc, "PANCREAS",
        d.get("pancreas_notes", "")
    )
    _section_label(
        doc, "SPLEEN",
        f"Spleen measuring {d.get('spleen_size', '')} cms."
    )
    _section_label(
        doc, "RIGHT KIDNEY",
        f"Measuring {d.get('right_kidney_size', '')} cms."
    )
    _section_label(
        doc, "LEFT KIDNEY",
        f"Measuring {d.get('left_kidney_size', '')} cms."
    )
    _section_label(
        doc, "URINARY BLADDER",
        d.get("urinary_bladder_notes", "Distended.")
    )
    _section_label(
        doc, "PROSTATE",
        d.get("prostate_notes", "")
    )
    _section_label(
        doc, "FREE FLUID",
        d.get("free_fluid_notes", "")
    )
    _section_label(
        doc, "BOWEL LOOPS",
        d.get("bowel_notes", "Visualized bowel loops show peristalsis.")
    )

    doc.add_paragraph()
    _heading(doc, "IMPRESSION:", size=11, center=False)
    doc.add_paragraph(d.get("impression", ""))

    doc.add_paragraph()
    _small_note(doc, STANDARD_NOTE)
    _doctor_signature(doc)

    return _save(doc, d.get("patient_name", "patient"), "abdomen_pelvis_male")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

SCAN_GENERATORS = {
    "early_pregnancy": generate_early_pregnancy_report,
    "nt_scan": generate_nt_scan_report,
    "anomaly_scan": generate_anomaly_scan_report,
    "growth_scan": generate_growth_scan_report,
    "follicular_study": generate_follicular_study_report,
    "abdomen_pelvis_female": generate_abdomen_pelvis_female_report,
    "abdomen_pelvis_male": generate_abdomen_pelvis_male_report,
}


def generate_report(scan_type: str, data: dict) -> Path:
    """
    Generate a report for the given scan type.

    Args:
        scan_type: One of the keys in SCAN_GENERATORS.
        data: Patient and measurement data dict.

    Returns:
        Path to the generated .docx file.
    """
    if scan_type not in SCAN_GENERATORS:
        raise ValueError(
            f"Unknown scan type '{scan_type}'. "
            f"Valid types: {list(SCAN_GENERATORS)}"
        )
    return SCAN_GENERATORS[scan_type](data)


# ---------------------------------------------------------------------------
# Interactive CLI
# ---------------------------------------------------------------------------

SCAN_MENU = {
    "1": ("early_pregnancy", "Early Pregnancy Scan"),
    "2": ("nt_scan", "NT Scan (Nuchal Translucency)"),
    "3": ("anomaly_scan", "Anomaly Scan"),
    "4": ("growth_scan", "Growth Scan"),
    "5": ("follicular_study", "Follicular Study"),
    "6": ("abdomen_pelvis_female", "Abdomen & Pelvis – Female"),
    "7": ("abdomen_pelvis_male", "Abdomen & Pelvis – Male"),
}

FIELD_PROMPTS: dict[str, list[tuple[str, str]]] = {
    "early_pregnancy": [
        ("patient_name", "Patient name"),
        ("age", "Age"),
        ("date", f"Date (default: {datetime.today().strftime('%d/%m/%Y')})"),
        ("ref_by", f"Referred by (default: {REFERRING_DEFAULT})"),
        ("lmp", "LMP"),
        ("clinical_details", "Clinical details"),
        ("crl", "CRL (cms)"),
        ("ga_scan_weeks", "GA by scan – weeks"),
        ("ga_scan_days", "GA by scan – days"),
        ("ga_lmp_weeks", "GA by LMP – weeks"),
        ("ga_lmp_days", "GA by LMP – days"),
        ("edd_scan", "EDD by scan"),
        ("edd_lmp", "EDD by LMP"),
        ("fhr", "FHR (BPM)"),
        ("cervix_length", "Cervix length (cms)"),
        ("impression", "Impression (leave blank — doctor fills in Word)"),
    ],
    "nt_scan": [
        ("patient_name", "Patient name"),
        ("age", "Age"),
        ("date", f"Date (default: {datetime.today().strftime('%d/%m/%Y')})"),
        ("ref_by", f"Referred by (default: {REFERRING_DEFAULT})"),
        ("lmp", "LMP"),
        ("ga_lmp_weeks", "GA from LMP – weeks"),
        ("ga_lmp_days", "GA from LMP – days"),
        ("edd_lmp", "EDD by LMP"),
        ("crl", "CRL (cms)"),
        ("aua_weeks", "AUA – weeks"),
        ("aua_days", "AUA – days"),
        ("fhr", "FHR (BPM)"),
        ("edd_scan", "EDD by scan"),
        ("nasal_bone", "Nasal bone (mm)"),
        ("nt", "Nuchal translucency (mm)"),
        ("ductus_venosus", "Ductus venosus (default: Normal a wave)"),
        ("liquor", "Liquor (default: Adequate)"),
        ("placenta_location", "Placenta location"),
        ("placenta_grade", "Placenta grade"),
        ("cervix_length", "Cervix length (cms)"),
        ("doppler_right_pi", "Right uterine artery PI"),
        ("doppler_left_pi", "Left uterine artery PI"),
        ("impression", "Impression (leave blank — doctor fills in Word)"),
    ],
    "anomaly_scan": [
        ("patient_name", "Patient name"),
        ("age", "Age"),
        ("date", f"Date (default: {datetime.today().strftime('%d/%m/%Y')})"),
        ("ref_by", f"Referred by (default: {REFERRING_DEFAULT})"),
        ("lmp", "LMP"),
        ("clinical_details", "Clinical details (e.g. G2P1 with 5MOA)"),
        ("foetal_lie", "Foetal lie/presentation (e.g. cephalic, transverse)"),
        ("bpd", "BPD (cms)"), ("hc", "HC (cms)"),
        ("fl", "FL (cms)"), ("ac", "AC (cms)"),
        ("hl", "HL (cms)"), ("tl", "TL (cms)"),
        ("rl", "RL (cms)"), ("fib", "FIB (cms)"), ("ul", "UL (cms)"),
        ("aua_weeks", "AUA – weeks"), ("aua_days", "AUA – days"),
        ("ga_lmp_weeks", "GA by LMP – weeks"), ("ga_lmp_days", "GA by LMP – days"),
        ("edd_scan", "EDD by scan"), ("edd_lmp", "EDD by LMP"),
        ("efw", "EFW (gms)"), ("efw_error", "EFW error margin (gms)"),
        ("fl_ac_ratio", "FL/AC ratio (%)"),
        ("nasal_bone_length", "Nasal bone length (mm)"),
        ("nuchal_fold", "Nuchal fold thickness (mm)"),
        ("tcd", "Trans cerebellar diameter (cm)"),
        ("cisterna_magna", "Cisterna magna (cm)"),
        ("lvta", "Lateral ventricular atrium (cm)"),
        ("foot_length", "Foot length (cms)"),
        ("fhr", "FHR (BPM)"),
        ("liquor", "Liquor (default: Adequate)"),
        ("placenta_location", "Placenta location"),
        ("placenta_grade", "Placenta grade"),
        ("cervix_length", "Cervix length (cms)"),
        ("doppler_right_pi", "Right uterine artery PI"),
        ("doppler_left_pi", "Left uterine artery PI"),
        ("impression", "Impression (leave blank — doctor fills in Word)"),
        ("comments", "Additional comments (optional)"),
    ],
    "growth_scan": [
        ("patient_name", "Patient name"),
        ("age", "Age"),
        ("date", f"Date (default: {datetime.today().strftime('%d/%m/%Y')})"),
        ("ref_by", f"Referred by (default: {REFERRING_DEFAULT})"),
        ("lmp", "LMP"),
        ("clinical_details", "Clinical details"),
        ("foetal_presentation", "Foetal presentation (default: cephalic)"),
        ("edd_lmp", "EDD by LMP"),
        ("bpd", "BPD (cms)"), ("hc", "HC (cms)"),
        ("fl", "FL (cms)"), ("ac", "AC (cms)"),
        ("aua_weeks", "AUA – weeks"), ("aua_days", "AUA – days"),
        ("ga_lmp_weeks", "GA by LMP – weeks"), ("ga_lmp_days", "GA by LMP – days"),
        ("edd_scan", "EDD by scan"),
        ("efw", "EFW (gms)"), ("efw_error", "EFW error margin (gms)"),
        ("fhr", "FHR (BPM)"),
        ("liquor", "Liquor (e.g. adequate)"),
        ("afi", "AFI (cms)"),
        ("placenta_location", "Placenta location"),
        ("placenta_grade", "Placenta grade"),
        ("cervix_length", "Cervix length (cms)"),
        ("doppler_right_pi", "Right uterine artery PI"),
        ("doppler_left_pi", "Left uterine artery PI"),
        ("doppler_umbilical_pi", "Umbilical artery PI"),
        ("doppler_mca_pi", "MCA flow PI"),
        ("impression", "Impression (leave blank — doctor fills in Word)"),
    ],
    "follicular_study": [
        ("patient_name", "Patient name"),
        ("age", "Age"),
        ("date", f"Date (default: {datetime.today().strftime('%d/%m/%Y')})"),
        ("ref_by", f"Referred by (default: {REFERRING_DEFAULT})"),
        ("uterus_size", "Uterus size (cms)"),
        ("right_ovary_size", "Right ovary size (cms)"),
        ("left_ovary_size", "Left ovary size (cms)"),
        ("bilateral_notes", "Bilateral ovary notes (optional)"),
        ("impression", "Impression"),
    ],
    "abdomen_pelvis_female": [
        ("patient_name", "Patient name"),
        ("age", "Age"),
        ("date", f"Date (default: {datetime.today().strftime('%d/%m/%Y')})"),
        ("ref_by", f"Referred by (default: {REFERRING_DEFAULT})"),
        ("clinical_data", "Clinical data"),
        ("liver_size", "Liver size (cms)"),
        ("spleen_size", "Spleen size (cms)"),
        ("right_kidney_size", "Right kidney size (cms)"),
        ("left_kidney_size", "Left kidney size (cms)"),
        ("uterus_size", "Uterus size (cms)"),
        ("endometrium_mm", "Endometrium thickness (mm)"),
        ("right_ovary_size", "Right ovary size (cms)"),
        ("left_ovary_size", "Left ovary size (cms)"),
        ("free_fluid_notes", "Free fluid (default: No free fluid noted)"),
        ("impression", "Impression"),
    ],
    "abdomen_pelvis_male": [
        ("patient_name", "Patient name"),
        ("age", "Age"),
        ("date", f"Date (default: {datetime.today().strftime('%d/%m/%Y')})"),
        ("ref_by", f"Referred by (default: {REFERRING_DEFAULT})"),
        ("clinical_data", "Clinical data"),
        ("liver_size", "Liver size (cms)"),
        ("spleen_size", "Spleen size (cms)"),
        ("right_kidney_size", "Right kidney size (cms)"),
        ("left_kidney_size", "Left kidney size (cms)"),
        ("prostate_notes", "Prostate notes (default: Normal in size and echogenicity)"),
        ("free_fluid_notes", "Free fluid (default: No free fluid noted)"),
        ("impression", "Impression"),
    ],
}


def _prompt_data(scan_key: str) -> dict:
    """Interactively prompt for patient data fields."""
    today = datetime.today().strftime("%d/%m/%Y")
    data: dict = {}
    prompts = FIELD_PROMPTS.get(scan_key, [])

    print("\n--- Enter patient details (press Enter to use default/leave blank) ---\n")
    for key, label in prompts:
        val = input(f"  {label}: ").strip()
        if not val:
            if "date" in key:
                val = today
            elif "ref_by" in key:
                val = REFERRING_DEFAULT
        data[key] = val

    # Follicular study: prompt for table rows
    if scan_key == "follicular_study":
        print("\n  Follicular tracking table (enter up to 12 rows, blank Day to stop):")
        rows = []
        for i in range(12):
            day = input(f"    Row {i+1} – Day: ").strip()
            if not day:
                rows.extend([{}] * (12 - i))
                break
            et = input(f"    Row {i+1} – Endometrial thickness: ").strip()
            rf = input(f"    Row {i+1} – Right ovary follicle: ").strip()
            lf = input(f"    Row {i+1} – Left ovary follicle: ").strip()
            ff = input(f"    Row {i+1} – Free fluid: ").strip()
            rows.append({
                "day": day,
                "endometrial_thickness": et,
                "right_follicle": rf,
                "left_follicle": lf,
                "free_fluid": ff,
            })
        data["follicle_rows"] = rows

    return data


def _run_interactive() -> None:
    print("\n========================================")
    print("   Ultrasound Report Generator (DOCX)   ")
    print("========================================")
    print("\nSelect report type:")
    for num, (_, label) in SCAN_MENU.items():
        print(f"  {num}. {label}")

    choice = input("\nEnter number (1–7): ").strip()
    if choice not in SCAN_MENU:
        print("Invalid choice. Exiting.")
        sys.exit(1)

    scan_key, scan_label = SCAN_MENU[choice]
    print(f"\n>>> Generating: {scan_label}")

    patient_data = _prompt_data(scan_key)
    output_path = generate_report(scan_key, patient_data)
    print(f"\n✓ Report saved to: {output_path.resolve()}")
    print("  Open the .docx file in Word to review and add doctor's impression if needed.")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate ultrasound report as editable .docx"
    )
    parser.add_argument(
        "--scan",
        choices=list(SCAN_GENERATORS),
        help="Scan type (skips interactive menu)",
    )
    args = parser.parse_args()

    if args.scan:
        # Non-interactive: prompt for fields only
        data = _prompt_data(args.scan)
        path = generate_report(args.scan, data)
        print(f"Report saved to: {path.resolve()}")
    else:
        _run_interactive()


if __name__ == "__main__":
    main()
